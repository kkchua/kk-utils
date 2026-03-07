"""
RAG Service - Unified Service Layer

Provides unified interface for RAG operations that:
- Stores metadata in PostgreSQL
- Stores chunks in ChromaDB
- Ensures consistency across both databases

Usage:
    from kk_utils.rag import RAGService
    
    service = RAGService()
    
    # Upload document (both PostgreSQL + ChromaDB)
    result = service.upload_document(
        file_path="/path/to/file.pdf",
        user_id="demo_user",
        classification="PUBLIC"
    )
    
    # Delete document (from both databases)
    service.delete_document(doc_id)
    
    # Search (from ChromaDB)
    results = service.search("query")
"""

from typing import Optional, Dict, Any, List
from pathlib import Path
import logging
import hashlib
from datetime import datetime
import uuid

logger = logging.getLogger(__name__)


class RAGService:
    """
    Unified RAG service layer.
    
    Handles both PostgreSQL (metadata) and ChromaDB (chunks).
    Ensures data consistency across both databases.
    """
    
    def __init__(
        self,
        database_url: Optional[str] = None,
        collection_name: Optional[str] = None,  # If None, reads from config
    ):
        """
        Initialize RAG service.
        
        Args:
            database_url: PostgreSQL connection URL (optional, reads from env if not provided)
            collection_name: ChromaDB collection name (optional, reads from config if not provided)
        """
        self.database_url = database_url
        
        # Read collection name from config if not provided
        if collection_name is None:
            try:
                from kk_utils.config_loader import ConfigLoader
                config = ConfigLoader()
                rag_config = config.get("rag", {})
                self.collection_name = rag_config.get("storage", {}).get("collection_name", "digital_me")
            except:
                self.collection_name = "digital_me"  # Fallback
        else:
            self.collection_name = collection_name
        
        self._db_session = None
        self._rag_engine = None
        
        logger.info(f"RAGService initialized: collection={self.collection_name}")
    
    def _get_db_session(self):
        """Get or create database session."""
        if self._db_session is None:
            try:
                from sqlalchemy import create_engine
                from sqlalchemy.orm import sessionmaker
                
                # Get database URL from env if not provided
                db_url = self.database_url or os.getenv(
                    "DATABASE_URL",
                    "sqlite:///./app.db"
                )
                
                engine = create_engine(db_url)
                SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
                self._db_session = SessionLocal()
                
                logger.debug("Database session created")
                
            except Exception as e:
                logger.warning(f"Failed to create database session: {e}")
                logger.warning("PostgreSQL features will be disabled")
                self._db_session = None
        
        return self._db_session
    
    def _get_rag_engine(self):
        """Get or create RAG engine."""
        if self._rag_engine is None:
            from kk_utils.rag import RAGEngine, RAGConfig
            
            config = RAGConfig(
                chunking={
                    "strategy": "sentence",
                    "chunk_size": 500,
                    "chunk_overlap": 50
                },
                retrieval={
                    "default_top_k": 5,
                    "min_confidence": 0.15
                }
            )
            
            # Use the same path as backend API
            backend_root = Path(__file__).resolve().parent.parent.parent.parent
            persist_dir = backend_root / "personal-assistant" / "backend" / "data" / "digital_me_rag"
            
            self._rag_engine = RAGEngine(
                collection_name=self.collection_name,
                persist_directory=str(persist_dir),
                config=config
            )
            
            logger.debug(f"RAG engine created: {self.collection_name}")
        
        return self._rag_engine
    
    def upload_document(
        self,
        file_path: str,
        user_id: str = "demo_user",
        classification: str = "PUBLIC",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Upload document to both PostgreSQL and ChromaDB.
        
        Args:
            file_path: Path to file to upload
            user_id: User ID for ownership
            classification: Document classification (PUBLIC, INTERNAL, etc.)
            metadata: Additional metadata
        
        Returns:
            Upload result with document info
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Uploading document: {file_path.name}")
        
        # Step 1: Extract text from file
        text = self._extract_text(file_path)
        
        # Step 2: Generate document ID and metadata
        doc_id = str(uuid.uuid4())
        file_hash = self._calculate_file_hash(file_path)
        
        doc_metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "file_type": file_path.suffix.lower(),
            "file_hash": file_hash,
            "user_id": user_id,
            "classification": classification,
            "upload_date": datetime.now().isoformat(),
            "source": "rag_service",
            **(metadata or {}),
        }
        
        # Step 3: Insert metadata to PostgreSQL
        db_result = self._insert_document_metadata(doc_id, doc_metadata)
        
        # Step 4: Add chunks to ChromaDB
        rag = self._get_rag_engine()
        chunk_result = rag.add_document(
            doc_id=doc_id,
            text=text,
            metadata=doc_metadata
        )
        
        logger.info(f"Document uploaded: {doc_id} ({chunk_result.get('chunks_added', 0)} chunks)")
        
        return {
            "success": True,
            "doc_id": doc_id,
            "filename": file_path.name,
            "chunks_added": chunk_result.get("chunks_added", 0),
            "total_chunks": chunk_result.get("total_chunks", 0),
            "database": db_result,
        }
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from file based on type."""
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == ".pdf":
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = "".join(page.get_text() for page in doc)
                doc.close()
                return text
            
            elif suffix == ".docx":
                from docx import Document
                doc = Document(file_path)
                return "\n".join(para.text for para in doc.paragraphs)
            
            elif suffix in [".txt", ".md"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
                
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def _insert_document_metadata(
        self,
        doc_id: str,
        metadata: Dict[str, Any],
    ) -> Optional[Dict[str, Any]]:
        """Insert document metadata to PostgreSQL."""
        session = self._get_db_session()
        if session is None:
            logger.warning("PostgreSQL not available, skipping metadata insert")
            return None
        
        try:
            # Import models
            from app.models.document import Document  # Assuming document model exists
            
            doc = Document(
                id=doc_id,
                filename=metadata["filename"],
                file_size=metadata["file_size"],
                file_type=metadata["file_type"],
                file_hash=metadata["file_hash"],
                user_id=metadata["user_id"],
                classification=metadata["classification"],
                metadata_json=metadata,
            )
            
            session.add(doc)
            session.commit()
            
            logger.debug(f"Document metadata inserted: {doc_id}")
            
            return {
                "success": True,
                "doc_id": doc_id,
            }
            
        except Exception as e:
            logger.error(f"Failed to insert document metadata: {e}")
            session.rollback()
            return {
                "success": False,
                "error": str(e),
            }
    
    def delete_document(self, doc_id: str) -> Dict[str, Any]:
        """
        Delete document from both PostgreSQL and ChromaDB.
        
        Args:
            doc_id: Document ID to delete
        
        Returns:
            Delete result
        """
        logger.info(f"Deleting document: {doc_id}")
        
        results = {
            "postgresql": False,
            "chromadb": False,
        }
        
        # Delete from PostgreSQL
        session = self._get_db_session()
        if session:
            try:
                from app.models.document import Document
                doc = session.query(Document).filter(Document.id == doc_id).first()
                if doc:
                    session.delete(doc)
                    session.commit()
                    results["postgresql"] = True
                    logger.debug(f"Deleted from PostgreSQL: {doc_id}")
            except Exception as e:
                logger.error(f"Failed to delete from PostgreSQL: {e}")
                session.rollback()
        
        # Delete from ChromaDB
        try:
            rag = self._get_rag_engine()
            # ChromaDB delete by ID
            rag.collection.delete(ids=[doc_id])
            results["chromadb"] = True
            logger.debug(f"Deleted from ChromaDB: {doc_id}")
        except Exception as e:
            logger.error(f"Failed to delete from ChromaDB: {e}")
        
        return {
            "success": results["chromadb"],  # Success if ChromaDB deleted
            "doc_id": doc_id,
            "results": results,
        }
    
    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        Search documents in ChromaDB.
        
        Args:
            query: Search query
            top_k: Number of results
        
        Returns:
            List of matching chunks
        """
        rag = self._get_rag_engine()
        result = rag.query(question=query, top_k=top_k)
        return result.chunks if result.chunks else []
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get RAG statistics."""
        rag = self._get_rag_engine()
        return rag.get_stats()
    
    def close(self):
        """Cleanup resources."""
        if self._db_session:
            self._db_session.close()
            self._db_session = None
        
        if self._rag_engine:
            self._rag_engine = None
        
        logger.debug("RAGService resources cleaned up")
